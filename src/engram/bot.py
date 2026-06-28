from __future__ import annotations

import asyncio
import logging
import os
import secrets
import shutil
import tempfile
import time
from dataclasses import dataclass, field
from datetime import datetime
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Awaitable, Callable

from anthropic import Anthropic
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Message, Update
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

from . import gitsafe, pending_store
from .config import DEFAULT_CATEGORY, ENTITY_TYPE_FOLDERS, Config, load_config
from .dedupe import find_semantic_duplicate
from .embeddings import OpenAIEmbedder, SemanticIndex, hybrid_search
from .entities import extract_entities, rebuild_entity_pages, upsert_entity_page
from .fetcher import fetch_urls
from .inbox import clear_pending, find_pending, move_note, preview
from .linker import answer_from_vault, enrich_note
from .note_writer import (
    CapturedMessage,
    append_to_note,
    extract_urls,
    merge_into_note,
    write_note,
)
from .pdf import extract_pdf_text
from .relinker import relink_note
from .vault import VaultIndex, load_note_body, scan_vault, search_vault
from .vision import ocr_image
from .whisper import transcribe

LOG_FORMAT = "%(asctime)s %(levelname)s %(name)s: %(message)s"
DEFAULT_LOG_PATH = Path.home() / ".engram.log"
DEFAULT_STATE_DIR = Path.home() / ".engram"

logging.basicConfig(format=LOG_FORMAT, level=logging.INFO)
log = logging.getLogger("engram")


def _attach_file_logging(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handler = RotatingFileHandler(
        path, maxBytes=5_000_000, backupCount=3, encoding="utf-8"
    )
    handler.setFormatter(logging.Formatter(LOG_FORMAT))
    handler.setLevel(logging.INFO)
    logging.getLogger().addHandler(handler)

VAULT_REFRESH_SECONDS = 600
MEDIA_GROUP_DEBOUNCE = 2.0
PENDING_TTL_SECONDS = 3600
SEARCH_RESULT_LIMIT = 8
TG_MESSAGE_CHAR_LIMIT = 4000
ASK_THREAD_TTL_SECONDS = 3600
ASK_THREAD_MAX_TURNS = 6


@dataclass
class _GroupBuffer:
    messages: list[Message] = field(default_factory=list)
    task: asyncio.Task | None = None


@dataclass
class _PendingCapture:
    messages: list[Message]
    extra_text: str = ""
    created_at: float = 0.0
    pending_files: list[Path] = field(default_factory=list)
    forward_info: dict | None = None


@dataclass
class _AskThread:
    turns: list[tuple[str, str]]  # (question, answer) pairs in chronological order
    created_at: float = 0.0


@dataclass
class _ReviewItem:
    path: Path
    created_at: float = 0.0


class BotState:
    def __init__(self, config: Config, state_dir: Path | None = None):
        self.config = config
        self.anthropic = Anthropic(api_key=config.anthropic_api_key)
        self._vault_index: VaultIndex = VaultIndex()
        self._vault_loaded_at: float = 0.0
        self._media_groups: dict[str, _GroupBuffer] = {}
        self._pending: dict[str, _PendingCapture] = {}
        self._last_capture: dict[int, Path] = {}  # chat_id -> last note path
        self._ask_threads: dict[tuple[int, int], _AskThread] = {}  # (chat_id, bot_msg_id) -> thread
        self._review_items: dict[str, _ReviewItem] = {}
        self._state_dir = state_dir
        embedder = (
            OpenAIEmbedder(api_key=config.openai_api_key)
            if config.openai_api_key
            else None
        )
        self._semantic_index = SemanticIndex(config.base_dir, embedder)

    @property
    def _pending_path(self) -> Path | None:
        return None if self._state_dir is None else self._state_dir / "pending.json"

    @property
    def _pending_files_dir(self) -> Path | None:
        return None if self._state_dir is None else self._state_dir / "pending_files"

    def stash_pending_files(self, token: str, files: list[Path]) -> list[Path]:
        """Copy temp files into a stable per-token dir so they survive restarts.

        No-op when persistence is disabled (no state_dir).
        """
        if not files or self._pending_files_dir is None:
            return list(files)
        dest_dir = self._pending_files_dir / token
        dest_dir.mkdir(parents=True, exist_ok=True)
        out: list[Path] = []
        for src in files:
            if not src.exists():
                continue
            dest = dest_dir / src.name
            try:
                src.replace(dest)
            except OSError:
                shutil.copy2(src, dest)
                src.unlink(missing_ok=True)
            out.append(dest)
        return out

    def _drop_pending_files(self, token: str) -> None:
        if self._pending_files_dir is None:
            return
        shutil.rmtree(self._pending_files_dir / token, ignore_errors=True)

    def persist_pending(self) -> None:
        path = self._pending_path
        if path is None:
            return
        try:
            pending_store.write(path, self._pending)
        except OSError:
            log.exception("Failed to persist pending captures to %s", path)

    def load_pending(self, bot) -> int:
        path = self._pending_path
        if path is None:
            return 0
        restored = pending_store.load(path, bot, _PendingCapture)
        self._pending.update(restored)
        if restored:
            log.info("Restored %d pending capture(s) from %s", len(restored), path)
        return len(restored)

    def _cleanup_pending(self) -> None:
        now = time.time()
        stale = [t for t, p in self._pending.items() if now - p.created_at > PENDING_TTL_SECONDS]
        if not stale:
            return
        for t in stale:
            self._pending.pop(t, None)
            self._drop_pending_files(t)
        self.persist_pending()

    def _cleanup_ask_threads(self) -> None:
        now = time.time()
        stale = [
            k for k, th in self._ask_threads.items()
            if now - th.created_at > ASK_THREAD_TTL_SECONDS
        ]
        for k in stale:
            self._ask_threads.pop(k, None)

    def _cleanup_review_items(self) -> None:
        now = time.time()
        stale = [
            t for t, it in self._review_items.items()
            if now - it.created_at > PENDING_TTL_SECONDS
        ]
        for t in stale:
            self._review_items.pop(t, None)

    def vault_index(self) -> VaultIndex:
        if time.time() - self._vault_loaded_at > VAULT_REFRESH_SECONDS:
            self.refresh_vault()
        return self._vault_index

    def invalidate_vault_cache(self) -> None:
        self._vault_loaded_at = 0.0

    def refresh_vault(self) -> None:
        log.info("Scanning vault: %s", self.config.base_dir)
        self._vault_index = scan_vault(self.config.base_dir)
        self._vault_loaded_at = time.time()
        log.info(
            "Indexed %d notes, %d tags",
            len(self._vault_index.titles),
            len(self._vault_index.tags),
        )
        if self._semantic_index.enabled:
            try:
                self._semantic_index.refresh()
            except Exception:
                log.exception("Semantic index refresh failed; continuing without it")


def _truncate(text: str, limit: int = TG_MESSAGE_CHAR_LIMIT) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def _obsidian_link(vault_name: str, path: Path, base_dir: Path) -> str | None:
    try:
        rel = path.relative_to(base_dir)
    except ValueError:
        return None
    from urllib.parse import quote
    return f"obsidian://open?vault={quote(vault_name)}&file={quote(rel.as_posix())}"


def _is_authorized(state: BotState, update: Update) -> bool:
    user = update.effective_user
    return bool(user and user.id in state.config.allowed_user_ids)


async def _download_photo(message: Message, target_dir: Path, n: int) -> Path:
    photo = message.photo[-1]
    target_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"{ts}-{n}.jpg"
    file = await photo.get_file()
    full_path = target_dir / filename
    await file.download_to_drive(full_path)
    return full_path


async def _download_document(message: Message) -> tuple[Path, str]:
    doc = message.document
    file = await doc.get_file()
    suffix = Path(doc.file_name or "file").suffix or ".bin"
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    Path(tmp_path).unlink()
    await file.download_to_drive(tmp_path)
    return Path(tmp_path), (doc.file_name or f"document{suffix}")


def _extract_forward_info(messages: list[Message]) -> dict | None:
    """Pull forward-origin info from PTB v22 Message.forward_origin (or older fallbacks).

    Returns a dict with keys: 'from' (display string), 'date' (datetime|None),
    'message_id' (int|None). Returns None if no message is a forward.
    """
    for m in messages:
        origin = getattr(m, "forward_origin", None)
        if origin is not None:
            display: str | None = None
            msg_id: int | None = None
            chat = getattr(origin, "chat", None)
            if chat is not None:
                username = getattr(chat, "username", None)
                title = getattr(chat, "title", None)
                display = f"@{username}" if username else (title or None)
                msg_id = getattr(origin, "message_id", None)
            else:
                user = getattr(origin, "sender_user", None)
                if user is not None:
                    username = getattr(user, "username", None)
                    full_name = getattr(user, "full_name", None) or getattr(user, "first_name", None)
                    display = f"@{username}" if username else full_name
                else:
                    display = getattr(origin, "sender_user_name", None)
            return {
                "from": display,
                "date": getattr(origin, "date", None),
                "message_id": msg_id,
            }
        # Fallback for older PTB or when forward_origin is absent
        fwd_chat = getattr(m, "forward_from_chat", None)
        fwd_user = getattr(m, "forward_from", None)
        if fwd_chat is None and fwd_user is None:
            continue
        if fwd_chat is not None:
            username = getattr(fwd_chat, "username", None)
            title = getattr(fwd_chat, "title", None)
            display = f"@{username}" if username else (title or None)
            return {
                "from": display,
                "date": getattr(m, "forward_date", None),
                "message_id": getattr(m, "forward_from_message_id", None),
            }
        if fwd_user is not None:
            username = getattr(fwd_user, "username", None)
            full_name = getattr(fwd_user, "full_name", None) or getattr(fwd_user, "first_name", None)
            display = f"@{username}" if username else full_name
            return {
                "from": display,
                "date": getattr(m, "forward_date", None),
                "message_id": None,
            }
    return None


async def _download_voice(message: Message) -> Path:
    voice = message.voice or message.audio
    file = await voice.get_file()
    suffix = ".ogg" if message.voice else (Path(voice.file_name or "audio.ogg").suffix or ".ogg")
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    Path(tmp_path).unlink()  # mkstemp creates the file; we want a fresh path
    await file.download_to_drive(tmp_path)
    return Path(tmp_path)


def _ocr_collected_images(state: BotState, image_paths: list[Path]) -> str:
    chunks: list[str] = []
    for p in image_paths:
        text = ocr_image(p, state.anthropic)
        if text:
            chunks.append(text)
    return "\n\n".join(chunks)


async def _process_messages(
    state: BotState,
    messages: list[Message],
    *,
    extra_text: str = "",
    redo: bool = False,
    delete_existing: Path | None = None,
    forced_folder: str | None = None,
    pending_files: list[Path] | None = None,
    forward_info: dict | None = None,
) -> tuple[Path, str]:
    text_parts: list[str] = []
    photo_messages: list[tuple[int, Message]] = []
    for i, m in enumerate(messages, 1):
        if m.photo:
            photo_messages.append((i, m))
        body = m.text or m.caption or ""
        if body:
            text_parts.append(body)

    if extra_text.strip():
        text_parts.append(extra_text.strip())

    raw_text = "\n\n".join(text_parts)
    urls = extract_urls(raw_text)

    pre_ocr_paths: list[Path] = []
    if photo_messages and not raw_text.strip():
        # OCR photos to a temp dir so we can route by content first
        with tempfile.TemporaryDirectory() as td:
            tmp_dir = Path(td)
            for i, m in photo_messages:
                p = await _download_photo(m, tmp_dir, i)
                pre_ocr_paths.append(p)
            ocr_text = await asyncio.to_thread(_ocr_collected_images, state, pre_ocr_paths)
        if ocr_text:
            raw_text = ocr_text
            urls = extract_urls(raw_text)

    if raw_text.strip():
        fetched = await asyncio.to_thread(fetch_urls, urls) if urls else {}
        enrichment = await asyncio.to_thread(
            enrich_note,
            raw_text,
            state.vault_index(),
            state.anthropic,
            fetched,
            state.config.categories,
            redo=redo,
        )
        folder = enrichment.folder
        title = enrichment.title or None
        body_text = enrichment.summary
        related, tags = enrichment.related, enrichment.tags
        source_type = enrichment.source_type
        review_pending = enrichment.confidence == "low"
    else:
        folder = DEFAULT_CATEGORY
        title = None
        body_text = ""
        related, tags = [], []
        source_type = "image" if photo_messages else "other"
        review_pending = False

    if forced_folder is not None:
        folder = forced_folder

    if delete_existing is not None and delete_existing.exists():
        log.info("Redo: deleting existing note %s", delete_existing)
        delete_existing.unlink()
        state.invalidate_vault_cache()

    duplicate = (
        None if redo else _find_duplicate(state.vault_index(), urls, title)
    )

    if duplicate is None and not redo and body_text.strip():
        duplicate = await asyncio.to_thread(
            find_semantic_duplicate,
            state._semantic_index,
            state.anthropic,
            title,
            body_text,
        )

    fwd_from = forward_info.get("from") if forward_info else None
    fwd_at = forward_info.get("date") if forward_info else None
    fwd_msg_id = forward_info.get("message_id") if forward_info else None

    if duplicate is not None:
        attachments_dir = duplicate.parent / "attachments"
        images: list[str] = []
        for i, m in photo_messages:
            full = await _download_photo(m, attachments_dir, i)
            images.append(f"attachments/{full.name}")
        attachments = _move_pending_files(pending_files or [], attachments_dir)
        captured = CapturedMessage(
            text=body_text,
            title=title,
            source_urls=urls,
            images=images,
            attachments=attachments,
            source_type=source_type,
            created=datetime.now(),
            forwarded_from=fwd_from,
            forwarded_at=fwd_at,
            original_message_id=fwd_msg_id,
            review_pending=review_pending,
        )
        # Karpathy-wiki merge: rewrite the canonical note to integrate the new
        # source instead of appending a dated block. Snapshot first so a bad
        # rewrite is one `git revert` away; merge_into_note falls back to a safe
        # append if the LLM rewrite would drop frontmatter or an attachment.
        gitsafe.snapshot(
            state.config.base_dir, f"engram: pre-merge {duplicate.stem}"
        )
        path, merged = await asyncio.to_thread(
            merge_into_note, duplicate, captured, state.anthropic
        )
        state.invalidate_vault_cache()
        await asyncio.to_thread(_update_entities, state, path.stem, body_text)
        verb = "merged" if merged else "appended"
        return path, f"{path.parent.name} ({verb})"

    target_dir = state.config.base_dir / folder
    attachments_dir = target_dir / "attachments"
    images = []
    for i, m in photo_messages:
        full = await _download_photo(m, attachments_dir, i)
        images.append(f"attachments/{full.name}")
    attachments = _move_pending_files(pending_files or [], attachments_dir)

    captured = CapturedMessage(
        text=body_text,
        title=title,
        source_urls=urls,
        images=images,
        attachments=attachments,
        source_type=source_type,
        created=datetime.now(),
        forwarded_from=fwd_from,
        forwarded_at=fwd_at,
        original_message_id=fwd_msg_id,
        review_pending=review_pending,
    )
    path = write_note(captured, target_dir, related=related, tags=tags)
    state.invalidate_vault_cache()
    await asyncio.to_thread(_update_entities, state, path.stem, body_text)
    return path, folder


def _update_entities(state: BotState, source_title: str, body: str) -> None:
    """Extract entities from a captured note and grow their typed wiki pages.

    Best-effort: never raises into the capture path.
    """
    if not body.strip():
        return
    try:
        for ent in extract_entities(state.anthropic, source_title, body):
            upsert_entity_page(state.config.base_dir, ent, source_title)
    except Exception:
        log.exception("Entity update failed for %s", source_title)


def _move_pending_files(srcs: list[Path], target_dir: Path) -> list[str]:
    if not srcs:
        return []
    target_dir.mkdir(parents=True, exist_ok=True)
    out: list[str] = []
    for src in srcs:
        if not src.exists():
            continue
        dest = target_dir / src.name
        n = 2
        while dest.exists():
            dest = target_dir / f"{src.stem} ({n}){src.suffix}"
            n += 1
        try:
            src.replace(dest)
        except OSError:
            # Cross-device fallback: copy + unlink
            import shutil
            shutil.copy2(src, dest)
            src.unlink(missing_ok=True)
        out.append(f"attachments/{dest.name}")
    return out


def _find_duplicate(index: VaultIndex, urls: list[str], title: str | None) -> Path | None:
    for url in urls:
        match = index.find_by_url(url.rstrip("/"))
        if match is not None:
            log.info("Duplicate by URL %s -> %s", url, match.name)
            return match
    if title:
        match = index.find_by_title(title)
        if match is not None:
            log.info("Duplicate by title %r -> %s", title, match.name)
            return match
    return None


def _folder_keyboard(token: str, categories: tuple[str, ...]) -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton(c, callback_data=f"f|{token}|{i}")]
        for i, c in enumerate(categories)
    ]
    return InlineKeyboardMarkup(rows)


def _review_keyboard(token: str, categories: tuple[str, ...]) -> InlineKeyboardMarkup:
    move_buttons = [
        InlineKeyboardButton(f"→ {c}", callback_data=f"r|{token}|mv|{i}")
        for i, c in enumerate(categories)
    ]
    move_rows: list[list[InlineKeyboardButton]] = []
    for i in range(0, len(move_buttons), 2):
        move_rows.append(move_buttons[i : i + 2])
    action_row = [
        InlineKeyboardButton("✅ Reviewed", callback_data=f"r|{token}|mark"),
        InlineKeyboardButton("⏭ Skip", callback_data=f"r|{token}|skip"),
        InlineKeyboardButton("🗑 Delete", callback_data=f"r|{token}|del"),
    ]
    return InlineKeyboardMarkup(move_rows + [action_row])


async def _ask_folder(
    state: BotState,
    send: Callable[..., Awaitable],
    messages: list[Message],
    *,
    extra_text: str = "",
    pending_files: list[Path] | None = None,
    forward_info: dict | None = None,
) -> str:
    state._cleanup_pending()
    token = secrets.token_urlsafe(8)
    if forward_info is None:
        forward_info = _extract_forward_info(messages)
    stashed = state.stash_pending_files(token, list(pending_files or []))
    state._pending[token] = _PendingCapture(
        messages=messages,
        extra_text=extra_text,
        created_at=time.time(),
        pending_files=stashed,
        forward_info=forward_info,
    )
    state.persist_pending()
    kb = _folder_keyboard(token, state.config.categories)
    await send("Choose a folder:", reply_markup=kb)
    return token


TEXT_EXTENSIONS = {
    ".md", ".markdown", ".txt", ".text", ".log", ".csv", ".tsv",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env",
    ".xml", ".html", ".htm", ".rst", ".org", ".tex",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".sh", ".bash", ".zsh",
    ".go", ".rs", ".java", ".c", ".h", ".cpp", ".hpp", ".rb",
    ".css", ".scss", ".sass", ".sql",
}

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
DOC_MIMES = ("application/msword", "application/x-msword")


def extract_docx_text(path: Path) -> str:
    from docx import Document  # imported lazily so tests can monkeypatch this module-level fn

    document = Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def extract_doc_text(path: Path) -> str:
    """Extract text from a legacy .doc via macOS textutil."""
    import subprocess

    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        timeout=30,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"textutil failed (rc={result.returncode}): {result.stderr.decode('utf-8', 'replace')}"
        )
    return result.stdout.decode("utf-8", "replace").strip()


def _classify_document(mime: str, name: str) -> str:
    """Return one of: 'pdf', 'docx', 'doc', 'text', 'unsupported'."""
    lower = name.lower()
    if mime == "application/pdf" or lower.endswith(".pdf"):
        return "pdf"
    if mime == DOCX_MIME or lower.endswith(".docx"):
        return "docx"
    if mime in DOC_MIMES or lower.endswith(".doc"):
        return "doc"
    if mime.startswith("text/"):
        return "text"
    ext = Path(lower).suffix
    if ext in TEXT_EXTENSIONS:
        return "text"
    return "unsupported"


async def _handle_document(state: BotState, update: Update, message: Message) -> None:
    doc = message.document
    mime = (doc.mime_type or "").lower()
    name = doc.file_name or ""
    kind = _classify_document(mime, name)
    if kind == "pdf":
        await _handle_pdf_document(state, update, message)
        return
    if kind == "docx":
        await _handle_word_document(state, update, message, fmt="docx")
        return
    if kind == "doc":
        await _handle_word_document(state, update, message, fmt="doc")
        return
    if kind == "text":
        await _handle_text_document(state, update, message)
        return
    await update.effective_chat.send_message(
        f"Unsupported document type: {mime or 'unknown'}. "
        "Supported: PDFs, Word (.docx/.doc), and plain-text formats (.md, .txt, .csv, .json, code, etc.)."
    )


async def _handle_pdf_document(state: BotState, update: Update, message: Message) -> None:
    await update.effective_chat.send_message("Extracting PDF text…")
    pdf_path, _ = await _download_document(message)
    try:
        text = await asyncio.to_thread(extract_pdf_text, pdf_path)
    except Exception:
        log.exception("PDF extraction failed")
        text = ""
    caption = message.caption or ""
    combined = "\n\n".join(t for t in (caption, text) if t.strip())
    if not combined.strip():
        await update.effective_chat.send_message(
            "Could not extract any text from PDF — capturing as attachment with default routing."
        )
    forward_info = _extract_forward_info([message])
    await _ask_folder(
        state,
        update.effective_chat.send_message,
        [],
        extra_text=combined,
        pending_files=[pdf_path],
        forward_info=forward_info,
    )


async def _handle_word_document(
    state: BotState, update: Update, message: Message, *, fmt: str
) -> None:
    await update.effective_chat.send_message(
        f"Extracting {fmt.upper()} text…"
    )
    file_path, _ = await _download_document(message)
    extractor = extract_docx_text if fmt == "docx" else extract_doc_text
    try:
        text = await asyncio.to_thread(extractor, file_path)
    except Exception:
        log.exception("%s extraction failed", fmt)
        text = ""
    caption = message.caption or ""
    combined = "\n\n".join(t for t in (caption, text) if t.strip())
    if not combined.strip():
        await update.effective_chat.send_message(
            f"Could not extract any text from {fmt.upper()} — capturing as attachment with default routing."
        )
    forward_info = _extract_forward_info([message])
    await _ask_folder(
        state,
        update.effective_chat.send_message,
        [],
        extra_text=combined,
        pending_files=[file_path],
        forward_info=forward_info,
    )


async def _handle_text_document(state: BotState, update: Update, message: Message) -> None:
    await update.effective_chat.send_message("Reading text file…")
    file_path, _ = await _download_document(message)
    try:
        text = await asyncio.to_thread(
            lambda: file_path.read_text(encoding="utf-8", errors="replace")
        )
    except Exception:
        log.exception("Text read failed")
        text = ""
    finally:
        file_path.unlink(missing_ok=True)
    caption = message.caption or ""
    combined = "\n\n".join(t for t in (caption, text) if t.strip())
    if not combined.strip():
        await update.effective_chat.send_message(
            "File was empty — nothing to capture."
        )
        return
    forward_info = _extract_forward_info([message])
    await _ask_folder(
        state,
        update.effective_chat.send_message,
        [],
        extra_text=combined,
        forward_info=forward_info,
    )


async def _handle_voice(state: BotState, update: Update, message: Message) -> None:
    if not state.config.openai_api_key:
        await update.effective_chat.send_message(
            "Voice transcription not configured. Set OPENAI_API_KEY in .env to enable."
        )
        return
    await update.effective_chat.send_message("Transcribing voice…")
    audio_path = await _download_voice(message)
    try:
        text = await asyncio.to_thread(transcribe, audio_path, state.config.openai_api_key)
    finally:
        audio_path.unlink(missing_ok=True)
    if not text:
        await update.effective_chat.send_message("Could not transcribe the voice message.")
        return
    await _ask_folder(state, update.effective_chat.send_message, [], extra_text=text)


async def _handle_single(state: BotState, update: Update, message: Message) -> None:
    await _ask_folder(state, update.effective_chat.send_message, [message])


async def _flush_group(
    state: BotState, group_id: str, chat_id: int, app: Application
) -> None:
    await asyncio.sleep(MEDIA_GROUP_DEBOUNCE)
    buf = state._media_groups.pop(group_id, None)
    if not buf or not buf.messages:
        return

    async def send(text: str, **kwargs):
        await app.bot.send_message(chat_id, text, **kwargs)

    await _ask_folder(state, send, buf.messages)


def make_handlers(state: BotState):
    async def on_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            await update.effective_chat.send_message("Unauthorized.")
            return
        message = update.effective_message
        if message is None:
            return

        # /ask follow-up: text reply to a prior bot /ask answer continues the thread.
        reply_target = getattr(message, "reply_to_message", None)
        body = (message.text or "").strip()
        if (
            reply_target is not None
            and body
            and not message.photo
            and not message.voice
            and not message.audio
            and not message.document
        ):
            chat_id = update.effective_chat.id if update.effective_chat else None
            target_id = getattr(reply_target, "message_id", None)
            if chat_id is not None and target_id is not None:
                state._cleanup_ask_threads()
                thread = state._ask_threads.get((chat_id, target_id))
                if thread is not None:
                    await _do_ask(update, body, prior_turns=list(thread.turns))
                    return

        if message.voice or message.audio:
            await _handle_voice(state, update, message)
            return

        if message.document:
            await _handle_document(state, update, message)
            return

        group_id = message.media_group_id
        if group_id:
            buf = state._media_groups.setdefault(group_id, _GroupBuffer())
            buf.messages.append(message)
            if buf.task is None or buf.task.done():
                buf.task = asyncio.create_task(
                    _flush_group(state, group_id, update.effective_chat.id, context.application)
                )
            return

        await _handle_single(state, update, message)

    async def on_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            await update.effective_chat.send_message("Unauthorized.")
            return
        cats = ", ".join(state.config.categories)
        await update.effective_chat.send_message(
            "Send text, links, photos, voice messages, PDFs, Word docs, or any plain-text "
            "file (.md, .txt, .csv, .json, code, etc.) — I'll auto-route them.\n\n"
            f"Categories: {cats}\n\n"
            "Commands:\n"
            "• /search <query> — find notes by title/tag/body\n"
            "• /ask <question> — answer from your vault (reply to chain follow-ups)\n"
            "• /inbox — list notes flagged for review\n"
            "• /review — walk pending notes one at a time\n"
            "• /relink [folder] — refresh related-links (last note, or a whole folder)\n"
            "• /undo — delete the last capture\n"
            "• /edit <new text> — replace the last capture's source\n"
            "• /redo (reply to a previous message) — regenerate with the higher-quality model\n"
            "• /refresh — reindex the vault"
        )

    async def on_refresh(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        state.refresh_vault()
        await update.effective_chat.send_message(
            f"Reindexed {len(state._vault_index.titles)} notes."
        )

    async def on_redo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        message = update.effective_message
        target = message.reply_to_message if message else None
        if target is None:
            await update.effective_chat.send_message(
                "Reply /redo to your original message (the one with the URL or text)."
            )
            return

        original_text = target.text or target.caption or ""
        original_urls = extract_urls(original_text)
        existing: Path | None = None
        for url in original_urls:
            existing = state.vault_index().find_by_url(url.rstrip("/"))
            if existing is not None:
                break

        try:
            path, folder = await _process_messages(
                state, [target], redo=True, delete_existing=existing
            )
            chat_id = update.effective_chat.id if update.effective_chat else None
            if chat_id is not None:
                state._last_capture[chat_id] = path
            await update.effective_chat.send_message(f"Re-saved to {folder}/{path.name}")
        except Exception as e:
            log.exception("Redo failed")
            await update.effective_chat.send_message(f"Error: {e}")

    async def on_folder_choice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        query = update.callback_query
        if query is None:
            return
        if not _is_authorized(state, update):
            await query.answer("Unauthorized.", show_alert=True)
            return
        parts = (query.data or "").split("|")
        if len(parts) != 3 or parts[0] != "f":
            await query.answer()
            return
        _, token, idx_s = parts
        pending = state._pending.pop(token, None)
        if pending is None:
            await query.answer("Expired — please resend.", show_alert=True)
            return
        state.persist_pending()
        try:
            chosen = state.config.categories[int(idx_s)]
        except (ValueError, IndexError):
            await query.answer("Invalid choice.", show_alert=True)
            return
        await query.answer()
        try:
            path, folder = await _process_messages(
                state,
                pending.messages,
                extra_text=pending.extra_text,
                forced_folder=chosen,
                pending_files=pending.pending_files,
                forward_info=pending.forward_info,
            )
            chat_id = update.effective_chat.id if update.effective_chat else None
            if chat_id is not None:
                state._last_capture[chat_id] = path
            await query.edit_message_text(f"Saved to {folder}/{path.name}")
        except Exception as e:
            log.exception("Folder choice processing failed")
            await query.edit_message_text(f"Error: {e}")
        finally:
            state._drop_pending_files(token)

    async def on_search(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        query = " ".join(context.args or []).strip()
        if not query:
            await update.effective_chat.send_message(
                "Usage: /search <query> — searches note titles, tags, and bodies."
            )
            return
        hits = await asyncio.to_thread(
            hybrid_search,
            state.config.base_dir,
            query,
            SEARCH_RESULT_LIMIT,
            state._semantic_index,
        )
        if not hits:
            await update.effective_chat.send_message(f"No matches for: {query}")
            return
        vault_name = state.config.base_dir.name
        lines: list[str] = [f"*{len(hits)} match(es) for `{query}`:*", ""]
        for h in hits:
            lines.append(f"• [[{h.title}]] _(in {h.category}, score {h.score})_")
            if h.snippet:
                lines.append(f"  {h.snippet}")
            link = _obsidian_link(vault_name, h.path, state.config.base_dir)
            if link:
                lines.append(f"  {link}")
            lines.append("")
        await update.effective_chat.send_message(
            _truncate("\n".join(lines)), disable_web_page_preview=True
        )

    async def _do_ask(
        update: Update,
        question: str,
        prior_turns: list[tuple[str, str]] | None = None,
    ) -> None:
        await update.effective_chat.send_message("Thinking…")
        try:
            result = await asyncio.to_thread(
                answer_from_vault,
                question,
                state.config.base_dir,
                state.anthropic,
                prior_turns=prior_turns,
                semantic_index=state._semantic_index,
            )
        except Exception as e:
            log.exception("/ask failed")
            await update.effective_chat.send_message(f"Error: {e}")
            return
        body = result.answer
        if result.sources:
            src_titles = " · ".join(f"[[{h.title}]]" for h in result.sources)
            body = f"{body}\n\n_Sources: {src_titles}_"
        sent = await update.effective_chat.send_message(
            _truncate(body), disable_web_page_preview=True
        )
        chat_id = update.effective_chat.id if update.effective_chat else None
        sent_id = getattr(sent, "message_id", None)
        if chat_id is None or sent_id is None:
            return
        state._cleanup_ask_threads()
        new_turns = list(prior_turns or []) + [(question, result.answer)]
        if len(new_turns) > ASK_THREAD_MAX_TURNS:
            new_turns = new_turns[-ASK_THREAD_MAX_TURNS:]
        state._ask_threads[(chat_id, sent_id)] = _AskThread(
            turns=new_turns, created_at=time.time()
        )

    async def on_ask(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        question = " ".join(context.args or []).strip()
        if not question:
            await update.effective_chat.send_message(
                "Usage: /ask <question> — answers from your vault notes. "
                "Reply to my answer to continue the thread."
            )
            return
        await _do_ask(update, question, prior_turns=None)

    async def on_undo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        chat_id = update.effective_chat.id if update.effective_chat else None
        path = state._last_capture.pop(chat_id, None) if chat_id is not None else None
        if path is None:
            await update.effective_chat.send_message("Nothing to undo.")
            return
        if not path.exists():
            await update.effective_chat.send_message(
                f"Last capture {path.name} no longer exists."
            )
            return
        try:
            path.unlink()
        except OSError as e:
            await update.effective_chat.send_message(f"Could not delete: {e}")
            return
        state.invalidate_vault_cache()
        await update.effective_chat.send_message(
            f"Deleted {path.parent.name}/{path.name}"
        )

    async def on_edit(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        new_text = " ".join(context.args or []).strip()
        if not new_text:
            await update.effective_chat.send_message(
                "Usage: /edit <new text> — replaces the source of the last capture and re-enriches it."
            )
            return
        chat_id = update.effective_chat.id if update.effective_chat else None
        old_path = state._last_capture.get(chat_id) if chat_id is not None else None
        if old_path is None or not old_path.exists():
            await update.effective_chat.send_message(
                "No recent capture to edit. Send a new message instead."
            )
            return
        old_folder = old_path.parent.name
        try:
            old_path.unlink()
        except OSError as e:
            await update.effective_chat.send_message(f"Could not remove old note: {e}")
            return
        state.invalidate_vault_cache()
        try:
            path, folder = await _process_messages(
                state, [], extra_text=new_text, forced_folder=old_folder
            )
            if chat_id is not None:
                state._last_capture[chat_id] = path
            await update.effective_chat.send_message(
                f"Edited → {folder}/{path.name}"
            )
        except Exception as e:
            log.exception("/edit failed")
            await update.effective_chat.send_message(f"Error: {e}")

    async def _send_review_card(
        send: Callable[..., Awaitable], path: Path
    ) -> None:
        state._cleanup_review_items()
        token = secrets.token_urlsafe(8)
        state._review_items[token] = _ReviewItem(path=path, created_at=time.time())
        preview_text = preview(path, max_lines=6)
        rel = path.parent.name if path.parent != state.config.base_dir else "/"
        header = f"📥 *Review:* {path.stem}\n_(in {rel})_"
        body = f"{header}\n\n{preview_text}" if preview_text else header
        kb = _review_keyboard(token, state.config.categories)
        await send(_truncate(body), reply_markup=kb, disable_web_page_preview=True)

    async def on_inbox(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        pending = find_pending(state.config.base_dir)
        if not pending:
            await update.effective_chat.send_message("Inbox is empty — no pending reviews.")
            return
        lines = [f"📥 *{len(pending)} note(s) pending review:*", ""]
        for p in pending[:10]:
            rel = p.parent.name if p.parent != state.config.base_dir else "/"
            lines.append(f"• [[{p.stem}]] _(in {rel})_")
        if len(pending) > 10:
            lines.append(f"_…and {len(pending) - 10} more_")
        lines.append("")
        lines.append("Run /review to walk through them.")
        await update.effective_chat.send_message(
            _truncate("\n".join(lines)), disable_web_page_preview=True
        )

    async def on_review(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        pending = find_pending(state.config.base_dir)
        if not pending:
            await update.effective_chat.send_message("Nothing to review.")
            return
        await _send_review_card(update.effective_chat.send_message, pending[0])

    async def on_review_choice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        query = update.callback_query
        if query is None:
            return
        if not _is_authorized(state, update):
            await query.answer("Unauthorized.", show_alert=True)
            return
        parts = (query.data or "").split("|")
        if len(parts) < 3 or parts[0] != "r":
            await query.answer()
            return
        token, action = parts[1], parts[2]
        item = state._review_items.pop(token, None)
        if item is None or not item.path.exists():
            await query.answer("Expired or note missing.", show_alert=True)
            try:
                await query.edit_message_text("Review item expired — run /review again.")
            except Exception:
                pass
            return
        await query.answer()
        path = item.path
        try:
            if action == "mark":
                clear_pending(path)
                state.invalidate_vault_cache()
                result_msg = f"✅ Marked reviewed: {path.stem}"
            elif action == "del":
                path.unlink()
                state.invalidate_vault_cache()
                result_msg = f"🗑 Deleted: {path.stem}"
            elif action == "skip":
                result_msg = f"⏭ Skipped: {path.stem}"
            elif action == "mv" and len(parts) == 4:
                try:
                    cat = state.config.categories[int(parts[3])]
                except (ValueError, IndexError):
                    await query.edit_message_text("Invalid category.")
                    return
                target_dir = state.config.base_dir / cat
                clear_pending(path)
                moved = move_note(path, target_dir)
                state.invalidate_vault_cache()
                result_msg = f"📂 Moved to {cat}/{moved.name}"
            else:
                await query.edit_message_text("Unknown action.")
                return
            await query.edit_message_text(result_msg)
        except Exception as e:
            log.exception("Review action failed")
            await query.edit_message_text(f"Error: {e}")
            return

        # Advance to the next pending item, if any.
        chat = update.effective_chat
        if chat is None:
            return
        remaining = find_pending(state.config.base_dir)
        # Filter the just-processed path in case of race.
        remaining = [p for p in remaining if p != path]
        if remaining:
            await _send_review_card(chat.send_message, remaining[0])
        else:
            await chat.send_message("✨ Inbox cleared — nothing left to review.")

    async def on_relink(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        if not state._semantic_index.enabled:
            await update.effective_chat.send_message(
                "Semantic index disabled (no OPENAI_API_KEY). Cannot relink."
            )
            return
        # Always refresh first so newly-written notes are part of the candidate set.
        await asyncio.to_thread(state._semantic_index.refresh)

        target_arg = " ".join(context.args or []).strip()
        if not target_arg:
            chat_id = update.effective_chat.id if update.effective_chat else None
            last = state._last_capture.get(chat_id) if chat_id is not None else None
            if last is None or not last.exists():
                await update.effective_chat.send_message(
                    "No recent capture in this chat. Try `/relink <folder>`."
                )
                return
            changed, titles = await asyncio.to_thread(
                relink_note, last, state._semantic_index
            )
            if changed:
                links = " · ".join(f"[[{t}]]" for t in titles) or "(none)"
                await update.effective_chat.send_message(
                    f"Relinked {last.stem}: {links}",
                    disable_web_page_preview=True,
                )
            else:
                await update.effective_chat.send_message(
                    f"No change for {last.stem}."
                )
            return

        # Folder mode.
        folder_dir = state.config.base_dir / target_arg
        if not folder_dir.is_dir():
            await update.effective_chat.send_message(
                f"Folder not found: {target_arg}"
            )
            return
        notes = [p for p in folder_dir.glob("*.md")]
        if not notes:
            await update.effective_chat.send_message(
                f"No notes in {target_arg}."
            )
            return
        await update.effective_chat.send_message(
            f"Relinking {len(notes)} note(s) in {target_arg}…"
        )
        changed_count = 0
        for p in notes:
            try:
                changed, _ = await asyncio.to_thread(
                    relink_note, p, state._semantic_index
                )
                if changed:
                    changed_count += 1
            except Exception:
                log.exception("relink failed for %s", p)
        state.invalidate_vault_cache()
        await update.effective_chat.send_message(
            f"Done. Updated {changed_count}/{len(notes)} note(s)."
        )

    async def on_rebuild(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not _is_authorized(state, update):
            return
        base_dir = state.config.base_dir
        await update.effective_chat.send_message(
            "Rebuilding vault: snapshotting to git, merging duplicates, "
            "rebuilding entity pages… this may take a minute."
        )
        committed = await asyncio.to_thread(
            gitsafe.snapshot, base_dir, "engram: pre-rebuild snapshot"
        )

        merges = 0
        merge_skipped = not state._semantic_index.enabled
        if merge_skipped:
            await update.effective_chat.send_message(
                "⚠️ Skipping the duplicate-merge phase: no OPENAI_API_KEY, so "
                "semantic similarity is unavailable. Only entity pages will be "
                "rebuilt. Add an embedding key to enable merging."
            )
        else:
            await asyncio.to_thread(state._semantic_index.refresh)
            note_paths = _vault_note_paths(base_dir)
            merges = await asyncio.to_thread(
                merge_duplicate_notes,
                base_dir,
                state._semantic_index,
                state.anthropic,
                note_paths,
            )

        notes = [
            (p.stem, load_note_body(p)) for p in _vault_note_paths(base_dir)
        ]
        entity_pages = await asyncio.to_thread(
            rebuild_entity_pages, base_dir, state.anthropic, notes
        )
        state.invalidate_vault_cache()
        await asyncio.to_thread(
            gitsafe.snapshot, base_dir, "engram: post-rebuild"
        )
        snap = "snapshot saved" if committed else "no snapshot (git unavailable or clean)"
        merge_line = (
            "merge phase SKIPPED (no embedding key)"
            if merge_skipped
            else f"merged {merges} duplicate note(s)"
        )
        await update.effective_chat.send_message(
            f"Rebuild done. {merge_line}; wrote {entity_pages} entity page(s). "
            f"Pre-rebuild {snap}; revert with "
            "`git -C <vault> reset --hard HEAD~` if it looks wrong."
        )

    return (
        on_message, on_start, on_refresh, on_redo, on_folder_choice,
        on_search, on_ask, on_undo, on_edit,
        on_inbox, on_review, on_review_choice,
        on_relink, on_rebuild,
    )


def _vault_note_paths(base_dir: Path) -> list[Path]:
    """All capture notes, excluding attachments and the entity (wiki) folders."""
    skip = {"attachments", *ENTITY_TYPE_FOLDERS.values()}
    out: list[Path] = []
    for p in base_dir.rglob("*.md"):
        if any(part in skip for part in p.relative_to(base_dir).parts):
            continue
        out.append(p)
    # Deterministic order so rebuilds and their git diffs are reproducible.
    return sorted(out)


def main() -> None:
    config = load_config()
    log_path = Path(os.environ.get("LOG_FILE", DEFAULT_LOG_PATH)).expanduser()
    _attach_file_logging(log_path)
    log.info("Logging to file: %s", log_path)
    state_dir = Path(os.environ.get("ENGRAM_STATE_DIR", DEFAULT_STATE_DIR)).expanduser()
    state = BotState(config, state_dir=state_dir)
    state.refresh_vault()

    app = Application.builder().token(config.telegram_token).build()
    state.load_pending(app.bot)
    (
        on_message, on_start, on_refresh, on_redo, on_folder_choice,
        on_search, on_ask, on_undo, on_edit,
        on_inbox, on_review, on_review_choice,
        on_relink, on_rebuild,
    ) = make_handlers(state)

    app.add_handler(CommandHandler("start", on_start))
    app.add_handler(CommandHandler("refresh", on_refresh))
    app.add_handler(CommandHandler("redo", on_redo))
    app.add_handler(CommandHandler("search", on_search))
    app.add_handler(CommandHandler("ask", on_ask))
    app.add_handler(CommandHandler("undo", on_undo))
    app.add_handler(CommandHandler("edit", on_edit))
    app.add_handler(CommandHandler("inbox", on_inbox))
    app.add_handler(CommandHandler("review", on_review))
    app.add_handler(CommandHandler("relink", on_relink))
    app.add_handler(CommandHandler("rebuild", on_rebuild))
    app.add_handler(CallbackQueryHandler(on_folder_choice, pattern=r"^f\|"))
    app.add_handler(CallbackQueryHandler(on_review_choice, pattern=r"^r\|"))
    app.add_handler(
        MessageHandler(
            (filters.TEXT & ~filters.COMMAND)
            | filters.PHOTO
            | filters.CAPTION
            | filters.VOICE
            | filters.AUDIO
            | filters.Document.ALL,
            on_message,
        )
    )

    log.info(
        "Bot starting; base dir: %s; categories: %s; whisper: %s",
        config.base_dir,
        config.categories,
        "enabled" if config.openai_api_key else "disabled",
    )
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
